from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from .db import attempts
from .state import ai


def _normalize_title(raw: str) -> str:
    s = (raw or "").strip()
    s = re.sub(r"[_\-]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip(" .-_")
    return s


def infer_paper_title(path: Path, extracted_text: str = "") -> str:
    ext = path.suffix.lower()
    fallback = _normalize_title(path.stem) or "Paper"

    # Prefer embedded document metadata when present.
    if ext == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            md = reader.metadata or {}
            t = _normalize_title(str(getattr(md, "title", "") or ""))
            if t and t.lower() not in {"untitled", "none", "nan"}:
                return t
        except Exception:
            pass
    if ext == ".docx":
        try:
            import docx

            d = docx.Document(str(path))
            t = _normalize_title(str(d.core_properties.title or ""))
            if t and t.lower() not in {"untitled", "none", "nan"}:
                return t
        except Exception:
            pass

    # Fallback to first meaningful heading line from extracted text.
    text = extracted_text or ""
    for line in (x.strip() for x in text.splitlines()):
        if not line:
            continue
        if len(line) < 4 or len(line) > 120:
            continue
        if re.match(r"^((q|question)\s*\d+|\d+[\)\.])", line, re.I):
            continue
        return _normalize_title(line)

    return fallback


def parse_questions(text: str) -> list[dict[str, Any]]:
    lines = [x.strip() for x in text.splitlines() if x.strip()]
    raw = [x for x in lines if re.match(r"^((q|question)\s*\d+|\d+[\)\.]\s+)", x, re.I)] or [
        f"Q{i+1}: Practice question" for i in range(8)
    ]
    out = []
    for i, q in enumerate(raw, 1):
        d = "easy" if i <= 3 else ("medium" if i <= 7 else "hard")
        out.append(
            {
                "qid": f"Q{i}",
                "question_text": q,
                "difficulty": d,
                "concept": "Core" if i <= 2 else f"Concept {((i - 1) % 4) + 1}",
                "correct_answer": "N/A",
                "max_marks": 5 if d == "easy" else (8 if d == "medium" else 12),
                "predicted_seconds": 300 if d == "easy" else (540 if d == "medium" else 780),
            }
        )
    return out


def extract_text_from_file(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if ext == ".pdf":
        try:
            from pypdf import PdfReader

            return "\n".join((p.extract_text() or "") for p in PdfReader(str(path)).pages)
        except Exception:
            return ""
    if ext == ".docx":
        try:
            import docx

            d = docx.Document(str(path))
            return "\n".join(p.text for p in d.paragraphs)
        except Exception:
            return ""
    return ""


def ai_questions_from_text(text: str, subject: str) -> list[dict[str, Any]]:
    if not ai or not text.strip():
        return []
    try:
        prompt = (
            f"Extract questions from this {subject} paper. Return strict JSON list: "
            '[{"qid":"Q1","question_text":"...","difficulty":"easy|medium|hard","concept":"...","correct_answer":"...","max_marks":5,"predicted_seconds":300}]\n'
            f"TEXT:\n{text[:12000]}"
        )
        resp = ai.responses.create(model="gpt-4.1-mini", input=prompt)
        raw = getattr(resp, "output_text", "")
        m = re.search(r"\[.*\]", raw, re.S)
        arr = json.loads(m.group(0)) if m else []
        out = []
        for i, q in enumerate(arr, 1):
            out.append(
                {
                    "qid": q.get("qid") or f"Q{i}",
                    "question_text": q.get("question_text", ""),
                    "difficulty": q.get("difficulty", "medium"),
                    "concept": q.get("concept", "Core"),
                    "correct_answer": q.get("correct_answer", "N/A"),
                    "max_marks": float(q.get("max_marks", 5)),
                    "predicted_seconds": int(q.get("predicted_seconds", 600)),
                }
            )
        return out
    except Exception:
        return []


def grade(question: dict[str, Any], answer: str) -> tuple[float, str]:
    max_marks = float(question.get("max_marks", 5))
    if not answer.strip():
        return 0.0, "Not attempted"
    if ai:
        try:
            prompt = (
                f"Grade answer out of {max_marks}. Q:{question.get('question_text','')} "
                f"Expected:{question.get('correct_answer','N/A')} Student:{answer}. "
                'Return JSON {"score":num,"feedback":"..."}'
            )
            resp = ai.responses.create(model="gpt-4.1-mini", input=prompt)
            text = getattr(resp, "output_text", "")
            payload = json.loads(re.search(r"\{.*\}", text, re.S).group(0))
            score = max(0.0, min(float(payload.get("score", 0)), max_marks))
            return score, str(payload.get("feedback", "AI graded"))
        except Exception:
            pass
    c = str(question.get("correct_answer", "")).lower().strip()
    if c and c != "n/a" and c in answer.lower():
        return max_marks, "Matched key answer"
    return round(max_marks * min(1.0, max(0.3, len(answer) / 120.0)), 2), "Heuristic grading"


def concept_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    bucket = {}
    for r in rows:
        if not r.get("attempted"):
            continue
        k = (r.get("subject", "General"), r.get("concept", "General"))
        c = bucket.setdefault(
            k,
            {"subject": k[0], "concept": k[1], "attempted": 0, "scored": 0.0, "possible": 0.0, "time": 0},
        )
        c["attempted"] += 1
        c["scored"] += float(r.get("mark_scored", 0))
        c["possible"] += float(r.get("max_marks", 0))
        c["time"] += int(r.get("time_spent_seconds", 0))
    return [
        {
            **x,
            "mastery_pct": round(100 * x["scored"] / max(1.0, x["possible"]), 2),
            "predicted_seconds": int(x["time"] / max(1, x["attempted"])),
        }
        for x in bucket.values()
    ]


def split_concepts(raw: str) -> list[str]:
    s = (raw or "").strip()
    if not s:
        return ["General"]
    # Keep "and" phrases intact, split only on hard separators.
    parts = [p.strip() for p in re.split(r"[,;|\n]+", s) if p.strip()]
    out = []
    seen = set()
    for p in parts or [s]:
        key = re.sub(r"\s+", " ", p).strip().lower()
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(re.sub(r"\s+", " ", p).strip())
    return out or ["General"]


def grouped_concept_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    bucket: dict[tuple[str, str], dict[str, Any]] = {}
    for r in rows:
        if not r.get("attempted"):
            continue
        subject = r.get("subject", "General")
        scored = float(r.get("mark_scored", 0))
        possible = float(r.get("max_marks", 0))
        spent = int(r.get("time_spent_seconds", 0))
        concepts = split_concepts(str(r.get("concept", "General")))
        for cpt in concepts:
            key = (subject, cpt)
            rec = bucket.setdefault(
                key,
                {"subject": subject, "concept": cpt, "attempted": 0, "scored": 0.0, "possible": 0.0, "time": 0},
            )
            rec["attempted"] += 1
            rec["scored"] += scored
            rec["possible"] += possible
            rec["time"] += spent
    out = []
    for x in bucket.values():
        out.append(
            {
                **x,
                "mastery_pct": round(100 * x["scored"] / max(1.0, x["possible"]), 2),
                "predicted_seconds": int(x["time"] / max(1, x["attempted"])),
            }
        )
    return out


def predict(user_id: str, subject: str):
    rows = [
        (datetime.fromisoformat(a["submitted_at"]), float(a.get("overall_pct", 0)))
        for a in attempts()
        if a.get("user_id") == user_id and a.get("subject") == subject and a.get("submitted_at")
    ]
    if len(rows) < 4:
        return None
    rows.sort(key=lambda x: x[0])
    if (rows[-1][0].date() - rows[0][0].date()).days < 21:
        return None
    h = len(rows) // 2
    a1 = sum(v for _, v in rows[:h]) / max(1, h)
    a2 = sum(v for _, v in rows[h:]) / max(1, len(rows) - h)
    return round(max(0.0, min(100.0, a2 + (a2 - a1) * 0.5)), 2)


def regression_hours_needed(user_id: str, subject: str, target_pct: float = 70.0):
    data: list[tuple[float, float]] = []
    for a in attempts():
        if a.get("user_id") != user_id or a.get("subject") != subject:
            continue
        x = float(a.get("duration_seconds", 0)) / 3600.0
        y = float(a.get("overall_pct", 0))
        if x > 0:
            data.append((x, y))

    if len(data) < 2:
        return None

    xs = [x for x, _ in data]
    ys = [y for _, y in data]
    x_mean = sum(xs) / len(xs)
    y_mean = sum(ys) / len(ys)

    var_x = sum((x - x_mean) ** 2 for x in xs)
    if var_x <= 1e-9:
        return None
    cov_xy = sum((x - x_mean) * (y - y_mean) for x, y in data)
    b = cov_xy / var_x
    a0 = y_mean - b * x_mean

    if b <= 1e-9:
        return None

    needed = (target_pct - a0) / b
    if needed < 0:
        needed = 0.0
    return round(min(needed, 40.0), 1)
